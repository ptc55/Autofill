import io
import random
import copy
from flask import Flask, render_template, request, send_file, session, flash, redirect, url_for
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


# --- 初始化 Flask 應用程式 ---
app = Flask(__name__)
app.secret_key = 'your-very-secret-and-random-key-final-version'

# --- 全局設定：班級列表 ---
CLASS_LIST = [
    "幼兒班第一組", "幼兒班第二組", "幼兒班第三組", "幼兒班第四組",
    "幼低班第一組", "幼低班第二組", "幼低班第三組", "幼低班第四組",
    "幼高班第一組", "幼高班第二組", "幼高班第三組", "幼高班第四組",
    "幼兒班第一組（下午）", "幼兒班第二組（下午）", "幼兒班第三組（下午）", "幼兒班第四組（下午）",
    "幼低班第一組（下午）", "幼低班第二組（下午）", "幼低班第三組（下午）", "幼低班第四組（下午）",
    "幼高班第一組（下午）", "幼高班第二組（下午）", "幼高班第三組（下午）", "幼高班第四組（下午）"
]

# --- 評語內容資料庫 ---
# ... 請將您現有的、完整的 comments_db 字典複製到這裡 ...
# 結構: { '副標題ID': {'title': '...', 'cell_location': {'table': T, 'row': R, 'col': C}, 'descriptions': [...]}}
# *** cell_location 座標已被全面修正 ***
comments_db = {
    # ----------- 認知發展 -----------
    'math_logic': {
        'title': '數理邏輯',
        # 修正: 'col' 從 1 改為 3
        'cell_location': {'table': 1, 'row': 0, 'col': 3},
        'descriptions': [
            {'text': '能辨認和點算數量',
             'sentences': ['能準確點算物件數量。', '對數與量的對應良好。', '已掌握基礎的點算技巧。', '能正確辨認具體數量。',
                           '在點算數量上表現穩定。']},
            {'text': '能比較物件的大小、長短',
             'sentences': ['能清楚比較物件的差異。', '對大小長短概念理解佳。', '可以準確區分物體屬性。',
                           '能運用比較詞彙描述。', '在比較任務中表現良好。']},
            {'text': '能進行簡單的分類',
             'sentences': ['能按指令將物品分類。', '已具備基礎的分類能力。', '能理解並執行分類任務。',
                           '可按單一屬性作分類。', '分類概念的掌握度高。']},
            {'text': '能認識基本平面圖形',
             'sentences': ['能辨認圓形、方形等圖形。', '對基本平面圖形有認識。', '可以說出常見圖形的名稱。',
                           '在圖形配對中表現不錯。', '已掌握數種基本圖形。']},
            {'text': '對數字和規律感興趣',
             'sentences': ['對數字遊戲顯得有興趣。', '樂於探索事物中的規律。', '在活動中能注意數理規律。',
                           '對數學概念抱持好奇心。', '喜愛操作與數理相關教具。']}
        ]
    },
    'problem_solving': {
        'title': '解難和創意思維',
        # 修正: 'col' 從 2 改為 3
        'cell_location': {'table': 1, 'row': 1, 'col': 3},
        'descriptions': [
            {'text': '遇到困難時願意嘗試',
             'sentences': ['面對挑戰時能積極嘗試。', '不畏困難，有嘗試的勇氣。', '遇到問題時願意動腦筋。',
                           '能主動尋求解決的方法。', '具備良好的解難態度。']},
            {'text': '能想出解決問題的方法',
             'sentences': ['能提出自己的解難方法。', '具備初步的問題解決能力。', '能想出簡單有效的方法。',
                           '在引導下能解決難題。', '解難的思路清晰且直接。']},
            {'text': '能發揮想像力進行創作',
             'sentences': ['在創作活動中富想像力。', '能用獨特的方式表達想法。', '作品中能見其豐富想像。',
                           '想像力豐富，想法新穎。', '樂於在創作中發揮巧思。']},
            {'text': '對新事物充滿好奇心',
             'sentences': ['對周遭新事物抱有好奇。', '探索慾強，喜愛接觸新知。', '能主動發問探索新事物。',
                           '對未知領域有探索精神。', '課堂中對新課題很投入。']},
            {'text': '能從不同角度思考問題',
             'sentences': ['能嘗試用不同方式解難。', '思路不固化，具靈活性。', '在引導下能轉換思考點。',
                           '具備初步的多角度思考。', '對問題有多元的看法。']}
        ]
    },
    'listening': {
        'title': '聆聽能力',
        # 修正: 'col' 從 2 改為 3
        'cell_location': {'table': 1, 'row': 2, 'col': 3},
        'descriptions': [
            {'text': '能安靜聆聽老師說話',
             'sentences': ['能專心聆聽老師的講解。', '課堂上能安靜聽老師說話。', '具備良好的課堂聆聽習慣。',
                           '聆聽時專注，不易分心。', '能耐心聽完老師的指令。']},
            {'text': '能安靜聆聽別人分享',
             'sentences': ['能尊重並聆聽同學發言。', '同儕分享時能安靜傾聽。', '具備輪流說話的意識。',
                           '團體活動中能聆聽他人。', '能耐心聽完同伴的分享。']},
            {'text': '能理解簡單的指令',
             'sentences': ['能理解並執行簡單指令。', '對老師的日常指令理解快。', '能準確完成單一步驟指令。',
                           '指令理解能力表現良好。', '在執行指令上頗為迅速。']},
            {'text': '能理解故事的主要內容',
             'sentences': ['聽故事後能說出大意。', '對故事內容的理解力好。', '能掌握故事的重點情節。',
                           '能回答與故事相關的問題。', '專心聽故事且能理解。']},
            {'text': '能記住聽到的訊息',
             'sentences': ['對聽過的內容記憶良好。', '能複述聽到的簡單訊息。', '記憶力佳，能記住細節。',
                           '聽覺記憶的表現不錯。', '能記住並傳達簡單訊息。']}
        ]
    },
    'speaking': {
        'title': '說話能力',
        # 修正: 'col' 從 2 改為 3
        'cell_location': {'table': 1, 'row': 3, 'col': 3},
        'descriptions': [
            {'text': '能運用完整句子說話',
             'sentences': ['能用完整句子表達意思。', '說話時句子結構趨向完整。', '已掌握運用完整句溝通。',
                           '能有條理地組織句子。', '樂於使用完整句子分享。']},
            {'text': '表達流暢，發音清晰',
             'sentences': ['說話表達流暢，口齒清晰。', '發音準確，表達清晰。', '語言表達清晰度和流暢度佳。',
                           '能清楚地表達自己的想法。', '說話有條理且發音清楚。']},
            {'text': '樂意在集體中發言',
             'sentences': ['樂於在團體面前分享。', '能勇敢地在集體中發言。', '在小組討論中能積極表達。',
                           '喜歡參與課堂問答環節。', '不怯場，敢於在人前說話。']},
            {'text': '能有禮貌地與人對話',
             'sentences': ['與人溝通時能使用禮貌語。', '懂得在對話中運用禮貌詞。', '已養成有禮貌的說話習慣。',
                           '能注意談話時的基本禮儀。', '與師長應對時態度有禮。']},
            {'text': '能描述日常生活的經歷',
             'sentences': ['能清楚描述自己的經歷。', '喜歡與人分享生活點滴。', '能有順序地講述事情。',
                           '能運用詞彙描述生活事件。', '在分享個人經驗時很投入。']}
        ]
    },
    'reading': {
        'title': '閱讀能力',
        # 修正: 'col' 從 2 改為 3
        'cell_location': {'table': 1, 'row': 4, 'col': 3},
        'descriptions': [
            {'text': '對圖書和文字感興趣',
             'sentences': ['喜愛翻閱圖書，探索內容。', '對文字符號抱有好奇心。', '對閱讀活動表現出興趣。',
                           '享受閱讀的樂趣與過程。', '能主動選擇感興趣的圖書。']},
            {'text': '能理解圖書內容',
             'sentences': ['能看懂圖畫並理解故事。', '能根據圖畫推測故事情節。', '對圖書內容有不錯的理解。',
                           '能明白書中角色的關係。', '閱讀後能說出書中大意。']},
            {'text': '有良好的閱讀習慣',
             'sentences': ['能安靜地閱讀，態度專注。', '懂得愛護圖書，輕輕翻閱。', '已建立初步的閱讀習慣。',
                           '會主動整理歸還圖書。', '閱讀時能保持正確姿勢。']},
            {'text': '能認讀簡單的常見字',
             'sentences': ['已能認讀一些常見單字。', '對學過的字詞記憶良好。', '在閱讀中能指認出單字。',
                           '識字量正在穩步增長中。', '對認讀自己的名字感興趣。']},
            {'text': '能理解文字的代表功能',
             'sentences': ['明白文字可用來記錄事情。', '知道文字符號帶有意義。', '對文字的功能有初步理解。',
                           '能理解書寫文字的作用。', '對環境中的文字感到好奇。']}
        ]
    },
    'writing': {
        'title': '書寫能力',
        # 修正: 'col' 從 2 改為 3
        'cell_location': {'table': 1, 'row': 5, 'col': 3},
        'descriptions': [
            {'text': '有正確的握筆姿勢',
             'sentences': ['握筆姿勢基本正確穩定。', '能以正確方式握持畫筆。', '在指導下能調整握筆手勢。',
                           '已掌握三指執筆的技巧。', '手眼協調，握筆有力。']},
            {'text': '能運用線條進行繪畫',
             'sentences': ['能運用不同線條來畫畫。', '繪畫時線條運用流暢。', '能控制線條表達創作意念。',
                           '喜歡用線條塗鴉和創作。', '畫作中的線條具表現力。']},
            {'text': '對書寫文字感興趣',
             'sentences': ['對模仿書寫文字有興趣。', '樂於嘗試書寫自己的名字。', '在畫作中會加入仿寫符號。',
                           '對寫字活動抱持正面態度。', '有書寫的意願和興趣。']},
            {'text': '能仿寫簡單的線條和字',
             'sentences': ['能跟隨虛線描寫簡單圖形。', '可仿寫直線、曲線等筆劃。', '能嘗試仿寫一些簡單字詞。',
                           '在描紅練習中表現穩定。', '具備初步的仿寫能力。']},
            {'text': '坐姿良好，專注書寫',
             'sentences': ['書寫時能保持良好坐姿。', '美勞活動時能專心投入。', '能安坐並完成書寫練習。',
                           '專注力佳，能持續創作。', '進行精細活動時很專心。']}
        ]
    },
    # ----------- 體能發展 -----------
    'gross_motor': {
        'title': '大肌肉的活動協調能力',
        # 無需修正，(1, 0) 是正確的合併儲存格起始位置
        'cell_location': {'table': 2, 'row': 0, 'col': 3},
        'descriptions': [
            {'text': '走路、跑步姿勢協調',
             'sentences': ['走、跑、跳等動作協調。', '跑步時身體平衡感良好。', '日常活動中動作敏捷。',
                           '能協調地完成基本動作。', '身體動作的協調性不錯。']},
            {'text': '能雙腳向前跳',
             'sentences': ['能雙腳離地向前跳躍。', '立定跳遠動作穩定流暢。', '跳躍動作的力量控制良好。',
                           '能輕鬆完成雙腳跳的動作。', '在跳躍遊戲中表現自如。']},
            {'text': '能單腳站立片刻',
             'sentences': ['能單腳站立保持平衡。', '身體平衡能力發展良好。', '在平衡遊戲中表現穩定。',
                           '單腳站立的持久性不錯。', '靜態平衡能力掌握得好。']},
            {'text': '能拍球及接球',
             'sentences': ['能連續拍打皮球數次。', '手眼協調，能接到大皮球。', '接拋球的動作反應迅速。',
                           '控球能力正在穩定進步。', '喜愛參與各類球類活動。']},
            {'text': '喜歡參與體能活動',
             'sentences': ['熱衷於參與戶外體能遊戲。', '享受大肌肉活動的樂趣。', '體能課上表現積極投入。',
                           '對各類體育活動很踴躍。', '喜愛跑動，精力充沛。']}
        ]
    },
    'fine_motor': {
        'title': '小肌肉的活動協調能力',
        # 修正: 'row' 從 2 改為 3 (評語區在標題下方一列)
        'cell_location': {'table': 2, 'row': 2, 'col': 3},
        'descriptions': [
            {'text': '能純熟地使用剪刀',
             'sentences': ['能沿直線和曲線剪紙。', '使用剪刀的技巧熟練。', '手眼協調，剪紙準確。', '能安全並有效地用剪刀。',
                           '小肌肉控制力佳，剪紙穩。']},
            {'text': '能穿珠子或繩子',
             'sentences': ['能完成穿珠子的精細動作。', '穿線動作迅速且準確。', '手部精細動作發展良好。',
                           '在穿線活動中表現出耐心。', '手眼協調能力足以穿繩。']},
            {'text': '能撕紙及搓紙糰',
             'sentences': ['能沿線條撕紙，控制力好。', '雙手協調，能搓出小紙球。', '手指靈活性及力量均不錯。',
                           '能完成撕、貼、搓等動作。', '在美勞活動中手部靈活。']},
            {'text': '能摺疊簡單的圖形',
             'sentences': ['能沿線對摺，摺疊整齊。', '已掌握基礎的摺紙技巧。', '能理解並跟隨摺紙步驟。',
                           '雙手協調，能完成摺工。', '對摺紙活動感到有興趣。']},
            {'text': '能扭開瓶蓋、扣鈕扣',
             'sentences': ['手腕轉動靈活，能開瓶蓋。', '能自行扣上或解開鈕扣。', '手指精細協調能力良好。',
                           '能處理生活中的精細操作。', '手部小肌肉發展成熟。']}
        ]
    },
    'self_care': {
        'title': '衞生習慣和自理能力',
        # 修正: 'row' 從 4 改為 5 (評語區在標題下方一列)
        'cell_location': {'table': 2, 'row': 4, 'col': 3},
        'descriptions': [
            {'text': '會自己上廁所及洗手',
             'sentences': ['能獨立如廁並清潔雙手。', '已養成飯前便後洗手習慣。', '自理能力佳，衛生習慣好。',
                           '個人衛生意識正在提高。', '能主動保持雙手的清潔。']},
            {'text': '會自己進食及收拾',
             'sentences': ['能獨立進食，無需協助。', '進食後會主動收拾餐具。', '自理能力強，用餐習慣好。',
                           '能熟練使用餐具並保持整潔。', '用餐後有收拾的責任感。']},
            {'text': '會保持個人儀容整潔',
             'sentences': ['懂得保持衣物和儀容整潔。', '有保持整潔的意識。', '會注意自己的儀容外觀。',
                           '能保持個人物品的整齊。', '在老師提醒下能整理儀容。']},
            {'text': '會自己穿脫衣服、鞋襪',
             'sentences': ['能自行穿脫 einfache 衣物。', '穿脫鞋襪的動作熟練。', '生活自理能力發展良好。',
                           '在日常自理方面表現獨立。', '已掌握穿脫衣物的技巧。']},
            {'text': '有良好的作息習慣',
             'sentences': ['作息規律，上課精神飽滿。', '有充足的睡眠，精神狀態好。', '正念時休時能安穩呼吸。',
                           '已建立健康的作息規律。', '生活作息穩定，適應園內生活。']}
        ]
    },
    # ----------- 情意及群性發展 -----------
    'self_image': {
        'title': '自我形象',
        # 修正: 'col' 從 2 改為 3
        'cell_location': {'table': 3, 'row': 0, 'col': 3},
        'descriptions': [
            {'text': '對自己有信心',
             'sentences': ['對自己的能力有信心。', '勇於接受挑戰，不怕失敗。', '在活動中表現自信大方。',
                           '自我肯定感強，態度積極。', '相信自己能完成任務。']},
            {'text': '能接納自己的長處短處',
             'sentences': ['能正面看待自己的優點。', '明白每個人都有長處短處。', '能坦然面對自己的不足。',
                           '對自我有客觀的認識。', '在引導下能接納自己。']},
            {'text': '認識自己的獨特性',
             'sentences': ['知道自己是獨一無二的。', '能說出自己的喜好和特點。', '對個人特質有初步認識。',
                           '能欣賞自己的獨特之處。', '在群體中能展現個性。']},
            {'text': '樂於嘗試新事物',
             'sentences': ['對新挑戰抱持開放態度。', '勇於嘗試未接觸過的事物。', '樂於參與各項新的活動。',
                           '探索精神可嘉，無畏嘗試。', '能積極投入新的學習中。']},
            {'text': '為自己的成功感到高興',
             'sentences': ['完成任務後有成就感。', '會因自己的進步感到自豪。', '懂得欣賞自己的努力成果。',
                           '能分享成功的喜悅。', '對個人成就抱持正面情感。']}
        ]
    },
    'self_management': {
        'title': '自我管理及表達感情的能力',
        # 修正: 'row' 從 1 改為 2, 'col' 從 1 改為 0
        'cell_location': {'table': 3, 'row': 1, 'col': 3},
        'descriptions': [
            {'text': '能適當表達個人情緒',
             'sentences': ['能用言語表達自己的感受。', '情緒表達方式大多是恰當的。', '開始學習管理及表達情緒。',
                           '能辨識並說出基本情緒。', '在引導下能適當表達情感。']},
            {'text': '能控制自己的行為和情緒',
             'sentences': ['情緒平穩，行為自控力佳。', '能遵守課堂的基本規則。', '衝動控制能力正在進步。',
                           '能調節自己的情緒反應。', '在群體中能管理好自己。']},
            {'text': '能面對挫折和失敗',
             'sentences': ['面對挫敗時情緒恢復快。', '能從失敗中學習並再嘗試。', '抗逆力不錯，能接受挑戰。',
                           '遇到困難時能保持冷靜。', '能以正面態度看待挫折。']},
            {'text': '能安靜地進行小組活動',
             'sentences': ['小組活動時能保持專注。', '能遵守小組活動的規則。', '能安靜投入於分配的任務。',
                           '在小組中能與人合作。', '自律性佳，能投入活動。']},
            {'text': '能明白別人的感受',
             'sentences': ['能察覺並關心同伴情緒。', '具備初步的同理心。', '能理解他人的喜怒哀樂。',
                           '開始學習站在他人角度想。', '對他人的感受具有敏感度。']}
        ]
    },
    'social_skills': {
        'title': '社交能力',
        # 修正: 'col' 從 2 改為 3
        'cell_location': {'table': 3, 'row': 4, 'col': 3},
        'descriptions': [
            {'text': '樂於與人交往',
             'sentences': ['性格開朗，喜愛與人互動。', '能主動結交朋友，融入群體。', '享受與同伴一起遊戲。',
                           '社交意願高，樂於合群。', '能自然地與師生交流。']},
            {'text': '能與人分享玩具和食物',
             'sentences': ['懂得與同伴分享物品。', '在遊戲中樂意分享玩具。', '分享意識良好，為人慷慨。',
                           '明白分享帶來的快樂。', '能主動提出與人分享。']},
            {'text': '能與人合作完成工作',
             'sentences': ['能與同伴合作進行遊戲。', '具備初步的團隊合作精神。', '在小組中能分工合作。',
                           '明白合作的重要性。', '樂於參與集體創作活動。']},
            {'text': '能解決與同伴的簡單紛爭',
             'sentences': ['能嘗試自行化解小紛爭。', '在引導下能協商解決問題。', '懂得用說話代替爭執。',
                           '社交解難能力正在發展。', '與同伴有矛盾時能求助。']},
            {'text': '能遵守遊戲規則',
             'sentences': ['能理解並遵守遊戲規則。', '有公平競爭的體育精神。', '在群體遊戲中能守秩序。',
                           '明白規則對遊戲的重要。', '能愉快地參與規則遊戲。']}
        ]
    },
    'responsibility': {
        'title': '責任感和公德心',
        # 修正: 'row' 從 5 改為 6, 'col' 從 1 改為 0
        'cell_location': {'table': 3, 'row': 5, 'col': 3},
        'descriptions': [
            {'text': '能完成自己應做的事',
             'sentences': ['有責任心，能完成分內事。', '對老師交託的任務很負責。', '能堅持完成自己的工作。',
                           '做事有始有終，態度認真。', '對自己的行為負責。']},
            {'text': '能愛護公物和環境',
             'sentences': ['懂得愛惜學校的公物。', '有保持環境清潔的意識。', '能做到物歸原位。', '會主動收拾及整理環境。',
                           '愛護公物，不隨意破壞。']},
            {'text': '能遵守校園規則',
             'sentences': ['能遵守學校的各項規則。', '常規良好，能遵守秩序。', '明白並能遵守集體規則。',
                           '在校園生活中表現守規。', '對建立的規則表示認同。']},
            {'text': '樂於幫助老師和同學',
             'sentences': ['有愛心，樂於助人。', '會主動關心和幫助同學。', '是老師得力的小幫手。', '熱心助人，富同情心。',
                           '在班級中常幫助有需要者。']},
            {'text': '能做好小組長工作',
             'sentences': ['擔任小組長時認真負責。', '能投入並完成小組長任務。', '對服務工作有責任感。',
                           '為能替集體服務感高興。', '能與其他小組長合作。']}
        ]
    },
    # ----------- 美感及文化發展 -----------
    'creativity_appreciation': {
        'title': '創作及欣賞事物的能力',
        # 修正: 'row' 從 0 改為 1, 'col' 從 1 改為 0
        'cell_location': {'table': 4, 'row': 0, 'col': 3},
        'descriptions': [
            {'text': '對顏色、線條、形狀敏感',
             'sentences': ['對視覺元素有敏銳觀察力。', '喜愛運用繽紛色彩創作。', '能欣賞事物的美感元素。',
                           '對色彩和構圖有直覺。', '能運用不同線條形狀表達。']},
            {'text': '樂於參與美藝活動',
             'sentences': ['對美藝創作活動十分投入。', '享受創作過程帶來的樂趣。', '能大膽地進行藝術創作。',
                           '在美勞課上表現積極。', '喜愛透過藝術表達自己。']},
            {'text': '能運用不同物料創作',
             'sentences': ['能嘗試用多樣物料創作。', '樂於探索不同媒介的效果。', '能發揮創意利用回收物料。',
                           '物料運用上具想像力。', '能掌握不同工具物料特性。']},
            {'text': '能表達對作品的感受',
             'sentences': ['能分享自己作品的創作意念。', '能簡單描述對藝術品的感受。', '樂於與人交流創作想法。',
                           '能用言語表達美感體驗。', '能欣賞並評論同學作品。']},
            {'text': '能投入音樂及律動活動',
             'sentences': ['享受音樂，能跟隨節奏擺動。', '在音樂活動中表現投入。', '對旋律和節奏反應良好。',
                           '樂於透過律動表達情感。', '喜愛唱歌及參與音樂遊戲。']}
        ]
    },
    'culture_appreciation': {
        'title': '認識及欣賞本身和其他民族的文化',
        # 修正: 'row' 從 3 改為 4, 'col' 從 1 改為 0
        'cell_location': {'table': 4, 'row': 3, 'col': 3},
        'descriptions': [
            {'text': '對節日慶祝活動感興趣',
             'sentences': ['對傳統節日活動興致高。', '樂於參與各項節慶活動。', '能感受節日的愉快氣氛。',
                           '對節日背後的故事好奇。', '喜歡學習不同節日的習俗。']},
            {'text': '能認識本地文化特色',
             'sentences': ['對身處的社區有初步認識。', '能說出香港的一些特色。', '對本土文化抱持興趣。',
                           '樂於了解社區的人和事。', '對本地文化有親切感。']},
            {'text': '能尊重不同文化',
             'sentences': ['能尊重不同國籍的同學。', '對世界各地的文化感好奇。', '明白不同地方有不同習俗。',
                           '能以開放態度接觸新文化。', '具備初步的文化尊重意識。']},
            {'text': '能參與傳統文化遊戲',
             'sentences': ['喜愛參與傳統的民間遊戲。', '能投入富文化色彩的活動。', '在文化活動中學習和成長。',
                           '對傳統玩意兒感到新奇。', '樂在其中地體驗文化活動。']},
            {'text': '能欣賞不同地方的藝術',
             'sentences': ['能欣賞不同風格的音樂美術。', '對多元的藝術形式感興趣。', '樂於接觸不同文化藝術品。',
                           '能感受不同文化的美。', '能從藝術中認識世界。']}
        ]
    }
}


# --- 輔助函式 ---
def apply_and_format_text(paragraph, text, is_header=False):
    p_format = paragraph.paragraph_format
    p_format.space_before = Pt(0)
    p_format.space_after = Pt(0)
    if is_header:
        p_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = paragraph.add_run(text)
    try:
        run.style = 'CommentFont'
    except KeyError:
        font = run.font
        font.name = 'KaiTi'
        font.east_asia = 'KaiTi'
        font.size = Pt(12)


def fill_doc_with_data(doc, record):
    student_name = record.get('student_name', '')
    class_name = record.get('class_name', '')
    teacher_name = record.get('teacher_name', '')

    header_table = doc.tables[0]

    cell_name = header_table.cell(0, 1)
    cell_name.text = ''
    apply_and_format_text(cell_name.paragraphs[0], student_name, is_header=True)

    cell_class = header_table.cell(1, 1)
    cell_class.text = ''
    apply_and_format_text(cell_class.paragraphs[0], class_name, is_header=True)

    cell_teacher = header_table.cell(1, 4)
    cell_teacher.text = ''
    apply_and_format_text(cell_teacher.paragraphs[0], teacher_name, is_header=True)

    all_tables = doc.tables
    selections = record.get('selections', {})
    for category_id, selected_indices in selections.items():
        if category_id in comments_db and selected_indices:
            category_data = comments_db[category_id]
            table_idx = category_data['cell_location']['table']
            row_idx = category_data['cell_location']['row']
            col_idx = category_data['cell_location']['col']

            final_comments = []
            for desc_index_str in selected_indices:
                desc_index = int(desc_index_str)
                description_obj = category_data['descriptions'][desc_index]
                chosen_sentence = random.choice(description_obj['sentences'])
                final_comments.append(chosen_sentence)

            if final_comments:
                target_cell = all_tables[table_idx].cell(row_idx, col_idx)
                target_cell.text = ''
                apply_and_format_text(target_cell.paragraphs[0], ' '.join(final_comments))


def merge_documents(docs):
    if not docs:
        return None

    final_doc = docs[0]

    for doc in docs[1:]:
        final_doc.add_page_break()
        for element in doc.element.body:
            final_doc.element.body.append(element)

    return final_doc


# --- 網頁路由 ---
@app.route('/')
def index():
    """ 渲染主頁面，並傳入班級列表和上次的老師姓名 """
    records = session.get('records', [])
    last_teacher_name = session.get('last_teacher_name', '')  # 從 session 獲取老師姓名
    return render_template('index.html', db=comments_db, record_count=len(records), class_list=CLASS_LIST,
                           last_teacher_name=last_teacher_name)


def _get_record_from_form(form):
    """ 從表單中提取資料並存儲老師姓名到 session """
    teacher_name = form.get('teacher_name', '')
    if teacher_name:
        session['last_teacher_name'] = teacher_name  # 儲存老師姓名到 session

    return {
        'student_name': form.get('student_name', ''),
        'class_name': form.get('class_name', ''),
        'teacher_name': teacher_name,
        'selections': {key: form.getlist(key) for key in comments_db.keys()}
    }


@app.route('/save', methods=['POST'])
def save_data():
    if 'records' not in session:
        session['records'] = []

    record = _get_record_from_form(request.form)

    session['records'].append(record)
    session.modified = True
    flash(f"「{record['student_name']}」的資料已儲存！目前共有 {len(session['records'])} 筆記錄。", 'success')
    return redirect(url_for('index'))


@app.route('/generate_single', methods=['POST'])
def generate_single():
    doc = Document('template.docx')
    record = _get_record_from_form(request.form)

    fill_doc_with_data(doc, record)
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    filename = f"{record['student_name'] or 'report'}_single.docx"
    return send_file(
        file_stream, as_attachment=True, download_name=filename,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )


@app.route('/generate_batch')
def generate_batch():
    records = session.get('records', [])
    if not records:
        flash("沒有已儲存的資料可供輸出。", 'warning')
        return redirect(url_for('index'))

    docs_to_merge = []
    for record in records:
        doc = Document('template.docx')
        fill_doc_with_data(doc, record)
        docs_to_merge.append(doc)

    final_doc = merge_documents(docs_to_merge)

    session.pop('records', None)

    file_stream = io.BytesIO()
    final_doc.save(file_stream)
    file_stream.seek(0)

    return send_file(
        file_stream, as_attachment=True, download_name='batch_report_final.docx',
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )


@app.route('/clear_session')
def clear_session():
    count = len(session.get('records', []))
    session.pop('records', None)
    session.pop('last_teacher_name', None)  # 同時清空老師姓名記憶
    flash(f"已成功清空 {count} 筆暫存資料及老師姓名記憶。", 'info')
    return redirect(url_for('index'))


if __name__ == '__main__':
    app.run(debug=True)